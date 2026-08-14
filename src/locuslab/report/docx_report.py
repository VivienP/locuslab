"""Build report.docx via python-docx.

RA/QA-readable presentation of the run: title, run summary, input documents,
artifact inventory, findings by severity and type, finding detail table,
known limitations, audit / provenance summary, and reviewer next steps.

Conservative tone enforced via assert_no_forbidden_language on every prose
string this module emits.
"""

from __future__ import annotations

import datetime as dt
from collections.abc import Sequence
from pathlib import Path
from typing import cast

import docx
from docx.document import Document as DocxDocument

from locuslab.models import (
    Claim,
    Document,
    EvidenceLink,
    Finding,
    FindingSeverity,
    Source,
)
from locuslab.report.language import assert_no_forbidden_language

# Fixed timestamp keeps visible document properties stable across reruns.
_FIXED_TIMESTAMP = dt.datetime(2026, 1, 1, 0, 0, 0)

_SEVERITY_ORDER: tuple[str, ...] = (
    FindingSeverity.CRITICAL.value,
    FindingSeverity.MAJOR.value,
    FindingSeverity.MINOR.value,
    FindingSeverity.INFORMATIONAL.value,
)

_REVIEWER_NEXT_STEPS: tuple[str, ...] = (
    "Open findings.xlsx in your spreadsheet tool and sort by severity.",
    "For each finding, read the evidence column and trace the affected "
    "object IDs in graph.jsonl to the originating span and document.",
    "Record adjudication outcomes in the reviewer / review_notes / "
    "resolution columns of findings.xlsx.",
    "Re-run locus verify after dossier remediation to refresh the report "
    "package on the same set of inputs.",
)


def _add_paragraph(doc: DocxDocument, text: str, where: str) -> None:
    assert_no_forbidden_language(text, where)
    doc.add_paragraph(text)


def _add_heading(doc: DocxDocument, text: str, level: int, where: str) -> None:
    assert_no_forbidden_language(text, where)
    doc.add_heading(text, level=level)


def _stabilise_core_properties(doc: DocxDocument) -> None:
    core = doc.core_properties
    core.author = "LocusLab"
    core.last_modified_by = "LocusLab"
    core.title = "LocusLab Evidence Trace Audit Report"
    core.created = _FIXED_TIMESTAMP
    core.modified = _FIXED_TIMESTAMP


def _short(value: str, length: int = 12) -> str:
    return value[:length]


_DRAFT_MARKER = "DRAFT — for internal review and adjudication"


def _write_title_block(doc: DocxDocument, dossier_path: str, run_id: str) -> None:
    """Phase 7B cover page: Heading 0 title + DRAFT marker + dossier + run_id + page break.

    The DRAFT marker is plain bold text, not a Word watermark — watermarks
    require section-property hacks that add zero buyer value beyond the
    visible marker. Run ID stays as the temporal anchor (deterministic);
    no wall-clock date is added.
    """
    _add_heading(doc, "LocusLab Evidence Trace Audit Report", level=0, where="docx.title")
    assert_no_forbidden_language(_DRAFT_MARKER, "docx.title.draft_marker")
    draft_paragraph = doc.add_paragraph()
    draft_run = draft_paragraph.add_run(_DRAFT_MARKER)
    draft_run.bold = True
    _add_paragraph(doc, f"Dossier: {dossier_path}", "docx.title.dossier")
    _add_paragraph(doc, f"Run ID: {run_id}", "docx.title.run_id")
    doc.add_page_break()  # type: ignore[no-untyped-call]


def _write_run_summary(
    doc: DocxDocument, run_id: str, artifact_counts: dict[str, int]
) -> None:
    _add_heading(doc, "Run summary", level=1, where="docx.run_summary.heading")
    _add_paragraph(
        doc,
        f"Run ID: {run_id}. This report summarises the artifacts produced by a "
        f"local LocusLab verify run on the supplied dossier.",
        "docx.run_summary.body",
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    header = table.rows[0].cells
    header[0].text = "Artifact count"
    header[1].text = "Value"
    for key in (
        "documents",
        "spans",
        "claims",
        "citations",
        "sources",
        "evidence_links",
        "findings",
        "graph_records",
    ):
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(artifact_counts.get(key, 0))


def _write_input_documents(doc: DocxDocument, documents: Sequence[Document]) -> None:
    _add_heading(doc, "Input documents", level=1, where="docx.input_documents.heading")
    if not documents:
        _add_paragraph(
            doc,
            "No input documents were ingested during this run.",
            "docx.input_documents.empty",
        )
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid"
    header = table.rows[0].cells
    for idx, col in enumerate(
        ("document_id", "kind", "path", "sha256 (short)", "parser", "warnings")
    ):
        header[idx].text = col
    for d in sorted(documents, key=lambda x: x.document_id):
        row = table.add_row().cells
        row[0].text = d.document_id
        row[1].text = d.kind.value
        row[2].text = d.path
        row[3].text = _short(d.sha256)
        row[4].text = d.parser
        row[5].text = ", ".join(sorted({w.code.value for w in d.parse_warnings})) or "-"


def _write_sources_block(doc: DocxDocument, sources: Sequence[Source]) -> None:
    """Phase 7B sources block: Heading 2 sub-section under Input documents.

    Lists bibliography sources resolved during ingestion with their
    availability status. When sources is empty (e.g. demo dossier without
    a bibliography), the heading is still written and the body carries
    the explicit non-finding line so the reviewer sees the gap.
    """
    _add_heading(doc, "Sources", level=2, where="docx.sources.heading")
    if not sources:
        _add_paragraph(
            doc,
            "No bibliography sources resolved for this dossier.",
            "docx.sources.empty",
        )
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    header = table.rows[0].cells
    for idx, col in enumerate(
        ("source_id", "citation_key", "path", "availability_status")
    ):
        header[idx].text = col
    for s in sorted(sources, key=lambda x: x.source_id):
        row = table.add_row().cells
        row[0].text = s.source_id
        row[1].text = s.citation_key or "-"
        row[2].text = s.path or "-"
        row[3].text = s.availability_status


def _write_artifact_inventory(
    doc: DocxDocument, source_artifact_hashes: dict[str, str]
) -> None:
    _add_heading(
        doc, "Artifact inventory", level=1, where="docx.artifact_inventory.heading"
    )
    _add_paragraph(
        doc,
        "Artifact hashes are quoted verbatim from audit_manifest.json "
        "(source_artifact_hashes).",
        "docx.artifact_inventory.body",
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    header = table.rows[0].cells
    header[0].text = "Artifact"
    header[1].text = "SHA-256 (short)"
    for name in sorted(source_artifact_hashes):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = _short(source_artifact_hashes[name])


def _write_findings_summary(doc: DocxDocument, findings: Sequence[Finding]) -> None:
    _add_heading(doc, "Findings summary", level=1, where="docx.findings_summary.heading")

    by_severity: dict[str, int] = {sev: 0 for sev in _SEVERITY_ORDER}
    by_type: dict[str, int] = {}
    for f in findings:
        by_severity[f.severity.value] += 1
        by_type[f.finding_type] = by_type.get(f.finding_type, 0) + 1

    _add_heading(doc, "By severity", level=2, where="docx.findings_summary.severity_heading")
    severity_table = doc.add_table(rows=1, cols=2)
    severity_table.style = "Light Grid"
    severity_table.rows[0].cells[0].text = "Severity"
    severity_table.rows[0].cells[1].text = "Count"
    for sev in _SEVERITY_ORDER:
        row = severity_table.add_row().cells
        row[0].text = sev
        row[1].text = str(by_severity[sev])

    _add_heading(doc, "By finding type", level=2, where="docx.findings_summary.type_heading")
    type_table = doc.add_table(rows=1, cols=2)
    type_table.style = "Light Grid"
    type_table.rows[0].cells[0].text = "Finding type"
    type_table.rows[0].cells[1].text = "Count"
    for ftype in sorted(by_type):
        row = type_table.add_row().cells
        row[0].text = ftype
        row[1].text = str(by_type[ftype])


_DOCX_FINDING_DETAIL_COLUMNS: tuple[str, ...] = (
    "eco_id",
    "severity",
    "finding_type",
    "checker_id",
    "affected_object_ids",
    "evidence",
    "remediation_hint",
    "adjudication_state",
)


def _write_finding_detail(doc: DocxDocument, findings: Sequence[Finding]) -> None:
    _add_heading(doc, "Finding detail", level=1, where="docx.finding_detail.heading")
    if not findings:
        _add_paragraph(doc, "No findings were produced for this run.", "docx.finding_detail.empty")
        return
    table = doc.add_table(rows=1, cols=len(_DOCX_FINDING_DETAIL_COLUMNS))
    table.style = "Light Grid"
    header = table.rows[0].cells
    for idx, col in enumerate(_DOCX_FINDING_DETAIL_COLUMNS):
        header[idx].text = col
    for f in sorted(findings, key=lambda x: x.eco_id):
        # Defence in depth: the Phase 3 checkers already exclude forbidden
        # language from evidence / remediation_hint, but this assertion
        # protects the report layer from a future regression.
        assert_no_forbidden_language(f.evidence, f"docx.finding_detail.{f.eco_id}.evidence")
        assert_no_forbidden_language(
            f.remediation_hint, f"docx.finding_detail.{f.eco_id}.remediation"
        )
        row = table.add_row().cells
        row[0].text = f.eco_id
        row[1].text = f.severity.value
        row[2].text = f.finding_type
        row[3].text = f.checker_id
        row[4].text = ";".join(f.affected_object_ids)
        row[5].text = f.evidence
        row[6].text = f.remediation_hint
        row[7].text = f.adjudication_state.value


def _write_known_limitations(doc: DocxDocument, known_limitations: Sequence[str]) -> None:
    _add_heading(doc, "Known limitations", level=1, where="docx.known_limitations.heading")
    for item in known_limitations:
        assert_no_forbidden_language(item, "docx.known_limitations.item")
        doc.add_paragraph(item, style="List Bullet")


def _write_audit_provenance(
    doc: DocxDocument,
    extraction_methods: Sequence[str],
    checker_ids: Sequence[str],
    linking_methods: Sequence[str],
) -> None:
    _add_heading(
        doc, "Audit and provenance summary", level=1, where="docx.audit_provenance.heading"
    )
    _add_paragraph(
        doc,
        "Extraction methods observed: " + (", ".join(extraction_methods) or "(none)"),
        "docx.audit_provenance.extraction",
    )
    _add_paragraph(
        doc,
        "Checker IDs observed: " + (", ".join(checker_ids) or "(none)"),
        "docx.audit_provenance.checkers",
    )
    _add_paragraph(
        doc,
        "Linking methods observed: " + (", ".join(linking_methods) or "(none)"),
        "docx.audit_provenance.linking",
    )


def _write_reviewer_next_steps(doc: DocxDocument) -> None:
    _add_heading(
        doc, "Reviewer next steps", level=1, where="docx.reviewer_next_steps.heading"
    )
    for step in _REVIEWER_NEXT_STEPS:
        assert_no_forbidden_language(step, "docx.reviewer_next_steps.step")
        doc.add_paragraph(step, style="List Number")


def write_report_docx(
    *,
    documents: Sequence[Document],
    claims: Sequence[Claim],
    evidence_links: Sequence[EvidenceLink],
    findings: Sequence[Finding],
    audit_manifest: dict[str, object],
    dossier_path: str,
    artifact_counts: dict[str, int],
    path: Path,
    sources: Sequence[Source] = (),
) -> None:
    """Write the buyer-facing report.docx artifact.

    Phase 7B re-introduced the `sources` parameter (dropped in the Phase 5
    reviewer Q-3 pass) so the docx can carry a per-source availability
    table under "Input documents". The default is the empty tuple so
    callers that have not yet been updated keep the prior behaviour.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    _stabilise_core_properties(doc)

    run_id = str(audit_manifest["run_id"])

    _write_title_block(doc, dossier_path, run_id)
    _write_run_summary(doc, run_id, artifact_counts)
    _write_input_documents(doc, documents)
    _write_sources_block(doc, sources)
    _write_artifact_inventory(
        doc, dict(cast(dict[str, str], audit_manifest["artifact_hashes"]))
    )
    _write_findings_summary(doc, findings)
    _write_finding_detail(doc, findings)
    _write_known_limitations(
        doc, list(cast(Sequence[str], audit_manifest["known_limitations"]))
    )
    _write_audit_provenance(
        doc,
        sorted({c.extraction_method for c in claims}),
        sorted({f.checker_id for f in findings}),
        sorted({link.linking_method for link in evidence_links}),
    )
    _write_reviewer_next_steps(doc)
    doc.save(str(path))
