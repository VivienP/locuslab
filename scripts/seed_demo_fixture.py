"""Regenerate `fixtures/demo_dossier/` with real DOCX, PDF, and XLSX files.

The committed fixtures should be small but real binary documents so that the
Phase 1b readers can produce content spans against them. This script is the
canonical way to refresh them. It runs offline and is deterministic.

Usage:

    python scripts/seed_demo_fixture.py

The script overwrites the three demo files in place. README.md and the
bibliography directory layout are preserved.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook

from locuslab.report.ooxml import canonicalise_ooxml

FIXTURE_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "demo_dossier"
_FIXED_TIMESTAMP = dt.datetime(2026, 1, 1, 0, 0, 0)


def write_cer_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.core_properties.author = "LocusLab"
    doc.core_properties.last_modified_by = "LocusLab"
    doc.core_properties.created = _FIXED_TIMESTAMP
    doc.core_properties.modified = _FIXED_TIMESTAMP
    doc.core_properties.title = "LocusLab Demo CER"
    doc.add_heading("Clinical Evaluation Report - DemoDevice X100", level=1)

    doc.add_heading("1. Device Description", level=2)
    doc.add_paragraph(
        "DemoDevice X100 is a Class IIa active medical device intended for "
        "continuous non-invasive monitoring in adult patients."
    )

    doc.add_heading("2. Clinical Performance", level=2)
    doc.add_paragraph(
        "The pivotal study (Smith et al., 2023) reported a primary endpoint "
        "response rate of 87.4% (95% CI: 82.1-91.6) at 12 months in 412 patients."
    )
    doc.add_paragraph(
        "Adverse event rate was 3.2% across the safety population (n=412)."
    )

    doc.add_heading("3. Benefit-Risk Conclusion", level=2)
    doc.add_paragraph(
        "Based on the available clinical evidence, the benefit-risk profile of "
        "DemoDevice X100 is considered acceptable for the intended use [1]."
    )

    doc.add_heading("Appendix A - Endpoint Summary Table", level=2)
    table = doc.add_table(rows=3, cols=3)
    headers = table.rows[0].cells
    headers[0].text = "Endpoint"
    headers[1].text = "Value"
    headers[2].text = "Source"
    row1 = table.rows[1].cells
    row1[0].text = "Response rate at 12 months"
    row1[1].text = "87.4%"
    row1[2].text = "Smith 2023"
    row2 = table.rows[2].cells
    row2[0].text = "Adverse event rate"
    row2[1].text = "3.2%"
    row2[2].text = "Smith 2023"

    doc.save(path)
    canonicalise_ooxml(path)


def write_gspr_xlsx(path: Path) -> None:
    wb = Workbook()
    wb.properties.creator = "LocusLab"
    wb.properties.lastModifiedBy = "LocusLab"
    wb.properties.created = _FIXED_TIMESTAMP
    wb.properties.modified = _FIXED_TIMESTAMP
    wb.properties.title = "LocusLab Demo GSPR Mapping"
    ws = wb.active
    assert ws is not None
    ws.title = "GSPR"
    ws.append(["GSPR_ID", "Requirement", "Applicable", "Evidence_Document", "Status"])
    ws.append(["GSPR-01", "Risk management process documented", "Yes", "CER.docx", "Met"])
    ws.append(["GSPR-02", "Clinical evaluation according to Annex XIV", "Yes", "CER.docx", "Met"])
    ws.append(["GSPR-03", "Post-market surveillance plan in place", "Yes", "PMS.docx", "Met"])
    ws.append(["GSPR-04", "Labeling complies with Annex I 23", "Yes", "Labeling.pdf", "Not Met"])
    ws.append(["GSPR-05", "Software lifecycle per IEC 62304", "Yes", "", ""])
    wb.save(path)
    canonicalise_ooxml(path)


def write_minimal_pdf(path: Path) -> None:
    pages_text = [
        (
            "Smith J. et al. (2023). Pivotal study of DemoDevice X100. "
            "J Med Devices 12(3):201-215. "
            "Primary endpoint response rate: 87.4 percent at 12 months."
        ),
        (
            "Adverse events were reported in 3.2 percent of the n=412 cohort. "
            "Confidence interval 95 percent: 82.1 to 91.6."
        ),
    ]

    buf = bytearray()
    offsets: list[int] = []

    def emit(s: str) -> None:
        buf.extend(s.encode("latin-1"))

    def emit_bytes(b: bytes) -> None:
        buf.extend(b)

    def add_object(body: str) -> None:
        offsets.append(len(buf))
        emit(f"{len(offsets)} 0 obj\n{body}\nendobj\n")

    emit("%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    page_object_numbers: list[int] = []
    content_object_numbers: list[int] = []

    catalog_obj_num = 1
    pages_obj_num = 2
    font_obj_num = 3

    add_object(f"<< /Type /Catalog /Pages {pages_obj_num} 0 R >>")

    add_object("__PAGES_PLACEHOLDER__")
    pages_placeholder_index = len(offsets) - 1

    add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>")

    next_obj = 4
    for text in pages_text:
        content = f"BT /F1 12 Tf 72 720 Td ({_escape_pdf_string(text)}) Tj ET"
        content_bytes = content.encode("latin-1")
        content_obj_body = f"<< /Length {len(content_bytes)} >>\nstream\n{content}\nendstream"
        add_object(content_obj_body)
        content_object_numbers.append(next_obj)
        next_obj += 1

        page_body = (
            f"<< /Type /Page /Parent {pages_obj_num} 0 R "
            f"/MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_obj_num} 0 R >> >> "
            f"/Contents {content_object_numbers[-1]} 0 R >>"
        )
        add_object(page_body)
        page_object_numbers.append(next_obj)
        next_obj += 1

    kids = " ".join(f"{n} 0 R" for n in page_object_numbers)
    pages_body = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_numbers)} >>"

    full = bytes(buf)
    placeholder_marker = b"__PAGES_PLACEHOLDER__"
    placeholder_pos = full.index(placeholder_marker)
    rebuilt = bytearray(full[:placeholder_pos])
    rebuilt.extend(pages_body.encode("latin-1"))
    after = full[placeholder_pos + len(placeholder_marker):]
    rebuilt.extend(after)

    delta = len(pages_body) - len(placeholder_marker)
    for i in range(pages_placeholder_index + 1, len(offsets)):
        offsets[i] += delta

    xref_offset = len(rebuilt)
    xref_lines = [f"xref\n0 {len(offsets) + 1}\n", "0000000000 65535 f \n"]
    for offset in offsets:
        xref_lines.append(f"{offset:010d} 00000 n \n")
    xref_block = "".join(xref_lines)
    trailer = (
        f"trailer\n<< /Size {len(offsets) + 1} /Root {catalog_obj_num} 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )

    out = bytes(rebuilt) + xref_block.encode("latin-1") + trailer.encode("latin-1")
    path.write_bytes(out)


def _escape_pdf_string(s: str) -> str:
    return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def main() -> None:
    if not FIXTURE_DIR.exists():
        raise SystemExit(f"Fixture directory not found: {FIXTURE_DIR}")

    write_cer_docx(FIXTURE_DIR / "CER.docx")
    write_gspr_xlsx(FIXTURE_DIR / "GSPR_mapping.xlsx")
    write_minimal_pdf(FIXTURE_DIR / "bibliography" / "source-study.pdf")

    print(f"Regenerated fixtures under {FIXTURE_DIR}")


if __name__ == "__main__":
    main()
