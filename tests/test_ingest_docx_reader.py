from pathlib import Path

import pytest
from docx import Document as DocxDocument

from locuslab.ingest.docx_reader import read_docx
from locuslab.models import ParserWarningCode, SpanLocationKind


def test_read_docx_extracts_paragraph_and_table_spans() -> None:
    result = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")

    assert result.parser.startswith("python-docx:")
    assert result.warnings == ()
    assert len(result.spans) > 0

    paragraph_kinds = {span.location.kind for span in result.spans}
    assert SpanLocationKind.PARAGRAPH in paragraph_kinds
    assert SpanLocationKind.TABLE_CELL in paragraph_kinds


def test_read_docx_skips_whitespace_only_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "blank.docx"
    doc = DocxDocument()
    doc.add_paragraph("Real content here.")
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    doc.add_paragraph("Tail content.")
    doc.save(path)

    result = read_docx(path, document_id="doc_blank")

    texts = [span.text for span in result.spans]
    assert texts == ["Real content here.", "Tail content."]


def test_read_docx_attaches_nearest_heading_as_section() -> None:
    result = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")

    body_paragraphs = [
        span
        for span in result.spans
        if span.location.kind == SpanLocationKind.PARAGRAPH
        and not span.text.startswith(("1.", "2.", "3.", "Appendix", "Clinical Evaluation"))
    ]
    assert body_paragraphs, "expected at least one body paragraph in CER fixture"
    for span in body_paragraphs:
        assert span.section is not None
        assert span.section.startswith(("1.", "2.", "3.", "Appendix"))


def test_read_docx_links_every_span_to_document_id() -> None:
    result = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")

    assert all(span.document_id == "doc_test_cer" for span in result.spans)
    assert all(span.text for span in result.spans)
    assert all(span.span_id.startswith("span_") for span in result.spans)


def test_read_docx_span_ids_are_stable_across_calls() -> None:
    first = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")
    second = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")

    assert tuple(s.span_id for s in first.spans) == tuple(s.span_id for s in second.spans)


def test_read_docx_table_cell_label_encodes_position() -> None:
    result = read_docx(Path("fixtures/demo_dossier/CER.docx"), document_id="doc_test_cer")

    cells = [s for s in result.spans if s.location.kind == SpanLocationKind.TABLE_CELL]
    assert cells
    for cell in cells:
        assert cell.location.label is not None
        assert cell.location.label.startswith("t0:r")


def test_read_docx_emits_unreadable_warning_for_non_docx_bytes(tmp_path: Path) -> None:
    bogus = tmp_path / "bogus.docx"
    bogus.write_bytes(b"not really a docx file")

    result = read_docx(bogus, document_id="doc_bogus")

    assert result.spans == ()
    codes = {w.code for w in result.warnings}
    assert ParserWarningCode.EXTRACTION_UNREADABLE_FILE in codes


def test_read_docx_emits_empty_document_warning_for_empty_body(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    DocxDocument().save(path)

    result = read_docx(path, document_id="doc_empty")

    codes = {w.code for w in result.warnings}
    assert ParserWarningCode.EXTRACTION_EMPTY_DOCUMENT in codes
    assert result.spans == ()


def test_read_docx_rejects_missing_file(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        read_docx(tmp_path / "nope.docx", document_id="doc_missing")
