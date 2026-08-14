"""Cover page and sources block in report.docx (offline)."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document as DocxRead

REPO_ROOT = Path(__file__).parent.parent
_SRC = REPO_ROOT / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

from locuslab.models import Source  # noqa: E402
from locuslab.pipeline import verify_dossier  # noqa: E402
from locuslab.report.docx_report import _DRAFT_MARKER, write_report_docx  # noqa: E402

DEMO_DOSSIER = REPO_ROOT / "fixtures" / "demo_dossier"


def _read_paragraphs(path: Path) -> list[str]:
    doc = DocxRead(str(path))
    return [p.text for p in doc.paragraphs]


def _read_headings(path: Path) -> list[tuple[int, str]]:
    """Return (level, text) for each heading-style paragraph."""
    doc = DocxRead(str(path))
    out: list[tuple[int, str]] = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style.startswith("Heading "):
            try:
                level = int(style.split(" ", 1)[1])
            except (IndexError, ValueError):
                level = 0
            out.append((level, p.text))
        elif style == "Title":
            out.append((0, p.text))
    return out


class TestCoverPage:
    @pytest.fixture()
    def docx_path(self, tmp_path: Path) -> Path:
        verify_dossier(dossier_dir=DEMO_DOSSIER, output_dir=tmp_path)
        return tmp_path / "report.docx"

    def test_docx_contains_draft_marker(self, docx_path: Path) -> None:
        paragraphs = _read_paragraphs(docx_path)
        assert any(_DRAFT_MARKER in p for p in paragraphs), (
            f"Expected DRAFT marker {_DRAFT_MARKER!r} in cover; "
            f"paragraphs were {paragraphs[:5]!r}"
        )

    def test_draft_marker_appears_before_run_summary(self, docx_path: Path) -> None:
        paragraphs = _read_paragraphs(docx_path)
        draft_idx = next(i for i, p in enumerate(paragraphs) if _DRAFT_MARKER in p)
        # Run summary heading shows up after the draft marker.
        run_summary_idx = next(
            (i for i, p in enumerate(paragraphs) if p == "Run summary"),
            None,
        )
        assert run_summary_idx is not None, "Run summary heading missing"
        assert draft_idx < run_summary_idx

    def test_title_heading_zero_present(self, docx_path: Path) -> None:
        headings = _read_headings(docx_path)
        title_match = [t for lvl, t in headings if lvl == 0]
        assert any(
            "LocusLab Evidence Trace Audit Report" in t for t in title_match
        ), f"Expected level-0 title; got headings {headings[:5]}"


class TestSourcesBlock:
    @pytest.fixture()
    def docx_path(self, tmp_path: Path) -> Path:
        verify_dossier(dossier_dir=DEMO_DOSSIER, output_dir=tmp_path)
        return tmp_path / "report.docx"

    def test_sources_heading_present_under_input_documents(self, docx_path: Path) -> None:
        headings = _read_headings(docx_path)
        idx_input = next(
            (i for i, (lvl, t) in enumerate(headings) if lvl == 1 and t == "Input documents"),
            None,
        )
        assert idx_input is not None, f"Input documents heading missing: {headings}"
        # Next heading must be the Sources subsection at level 2.
        next_heading = headings[idx_input + 1] if idx_input + 1 < len(headings) else None
        assert next_heading == (2, "Sources"), (
            f"Expected Heading 2 'Sources' after Input documents; got {next_heading}"
        )

    def test_sources_table_has_four_columns_and_three_rows_for_demo(
        self, docx_path: Path
    ) -> None:
        doc = DocxRead(str(docx_path))
        # Demo dossier resolves 3 sources (well-known baseline).
        sources_tables = [
            t
            for t in doc.tables
            if t.rows and t.rows[0].cells[0].text == "source_id"
        ]
        assert len(sources_tables) == 1, (
            f"Expected exactly 1 sources table; found {len(sources_tables)}"
        )
        sources_table = sources_tables[0]
        assert len(sources_table.columns) == 4
        assert len(sources_table.rows) == 4  # header + 3 source rows
        header = [c.text for c in sources_table.rows[0].cells]
        assert header == ["source_id", "citation_key", "path", "availability_status"]


class TestSourcesBlockEmpty:
    def test_sources_empty_shows_no_resolved_line(self, tmp_path: Path) -> None:
        """When sources is empty, the Heading 2 + the explicit
        no-bibliography-resolved line must both be present."""
        write_report_docx(
            documents=(),
            claims=(),
            evidence_links=(),
            findings=(),
            audit_manifest={
                "run_id": "run_test_0",
                "artifact_hashes": {},
                "known_limitations": [],
            },
            dossier_path="/fake",
            artifact_counts={},
            path=tmp_path / "report.docx",
            sources=(),
        )
        paragraphs = _read_paragraphs(tmp_path / "report.docx")
        assert any(
            "No bibliography sources resolved for this dossier." in p for p in paragraphs
        ), f"Empty-sources sentinel paragraph missing; got {paragraphs[-10:]}"

    def test_sources_table_absent_when_empty(self, tmp_path: Path) -> None:
        write_report_docx(
            documents=(),
            claims=(),
            evidence_links=(),
            findings=(),
            audit_manifest={
                "run_id": "run_test_1",
                "artifact_hashes": {},
                "known_limitations": [],
            },
            dossier_path="/fake",
            artifact_counts={},
            path=tmp_path / "report.docx",
            sources=(),
        )
        doc = DocxRead(str(tmp_path / "report.docx"))
        # No sources table should be present when sources is empty.
        sources_tables = [
            t for t in doc.tables
            if t.rows and t.rows[0].cells[0].text == "source_id"
        ]
        assert len(sources_tables) == 0


class TestSourcesBlockExplicit:
    def test_explicit_sources_render_in_table(self, tmp_path: Path) -> None:
        sources = (
            Source(
                source_id="src_a",
                path="bibliography/a.pdf",
                citation_key="Smith2024",
                availability_status="local_fulltext",
            ),
            Source(
                source_id="src_b",
                path=None,
                citation_key="Jones2023",
                availability_status="unresolved",
            ),
        )
        write_report_docx(
            documents=(),
            claims=(),
            evidence_links=(),
            findings=(),
            audit_manifest={
                "run_id": "run_test_explicit",
                "artifact_hashes": {},
                "known_limitations": [],
            },
            dossier_path="/fake",
            artifact_counts={},
            path=tmp_path / "report.docx",
            sources=sources,
        )
        doc = DocxRead(str(tmp_path / "report.docx"))
        sources_tables = [
            t for t in doc.tables
            if t.rows and t.rows[0].cells[0].text == "source_id"
        ]
        assert len(sources_tables) == 1
        sources_table = sources_tables[0]
        # Header + 2 source rows.
        assert len(sources_table.rows) == 3
        row1 = [c.text for c in sources_table.rows[1].cells]
        row2 = [c.text for c in sources_table.rows[2].cells]
        assert row1 == ["src_a", "Smith2024", "bibliography/a.pdf", "local_fulltext"]
        # Source with no path uses "-" placeholder.
        assert row2 == ["src_b", "Jones2023", "-", "unresolved"]


class TestPipelineFindingsUnchanged:
    def test_findings_jsonl_unchanged_by_docx_sources_block(self, tmp_path: Path) -> None:
        """The sources block is docx-only; findings.jsonl must be byte-equal across
        runs (also confirmed by graph/audit determinism tests)."""
        a = tmp_path / "run_a"
        b = tmp_path / "run_b"
        verify_dossier(dossier_dir=DEMO_DOSSIER, output_dir=a)
        verify_dossier(dossier_dir=DEMO_DOSSIER, output_dir=b)
        assert (a / "findings.jsonl").read_bytes() == (b / "findings.jsonl").read_bytes()
