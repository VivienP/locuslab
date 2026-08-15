"""Packaged SSCP guidance assets, independent of repo cwd."""

from __future__ import annotations

import json
from importlib.resources import files
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from locuslab.pipeline import verify_dossier

REPO_ROOT = Path(__file__).resolve().parent.parent

PACKAGED_GUIDANCE_FILES = (
    "docs/THIRD_PARTY.md",
    "docs/rules/guidance/sscp/rule_pack.json",
    "docs/guidance/source_inventory.json",
    "docs/guidance/sources/mdcg/md_mdcg_2019_9_sscp_en.md",
    "docs/guidance/sources/mdcg/mdcg_2022-9_en.md",
    "docs/guidance/sources/eurlex/article_32.md",
    "docs/guidance/sources/eurlex/article_61.md",
)


def test_byte_compared_guidance_files_are_pinned_to_lf() -> None:
    attributes = (REPO_ROOT / ".gitattributes").read_text(encoding="utf-8")
    for pattern in (
        "docs/THIRD_PARTY.md text eol=lf",
        "docs/guidance/*.json text eol=lf",
        "docs/rules/guidance/**/*.json text eol=lf",
        "src/locuslab/resources/docs/**/*.json text eol=lf",
    ):
        assert pattern in attributes.splitlines()


def test_packaged_guidance_files_exist() -> None:
    root = files("locuslab.resources")
    for rel in PACKAGED_GUIDANCE_FILES:
        trav = root.joinpath(rel)
        assert trav.is_file(), f"missing packaged guidance file: {rel}"


def test_packaged_guidance_bytes_match_docs_source_of_truth() -> None:
    root = files("locuslab.resources")
    for rel in PACKAGED_GUIDANCE_FILES:
        packaged = root.joinpath(rel).read_bytes()
        canonical_relative = (
            Path("docs/THIRD_PARTY.md")
            if rel == "docs/THIRD_PARTY.md"
            else Path(rel)
        )
        canonical = (REPO_ROOT / canonical_relative).read_bytes()
        assert packaged == canonical, f"packaged copy drifted from {rel}"


def _write_mini_sscp_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Summary of safety and clinical performance")
    doc.add_paragraph("Intended purpose: synthetic mini SSCP for packaging tests.")
    doc.add_paragraph("Device description: packaging fixture only.")
    doc.save(str(path))


def test_verify_sscp_guidance_without_repo_cwd(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Guidance review must work when cwd is not the git checkout."""
    elsewhere = tmp_path / "not_the_repo"
    elsewhere.mkdir()
    dossier = tmp_path / "dossier"
    dossier.mkdir()
    _write_mini_sscp_docx(dossier / "synthetic_sscp_mini.docx")
    out = tmp_path / "out"
    monkeypatch.chdir(elsewhere)

    result = verify_dossier(dossier_dir=dossier, output_dir=out)

    assert result.n_guidance_review_items == 10
    review_path = out / "guidance_review.json"
    assert review_path.is_file()
    data = json.loads(review_path.read_text(encoding="utf-8"))
    assert data["n_review_items"] == 10
    assert (out / "guidance_review.md").is_file()
