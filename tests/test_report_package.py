"""report.docx, findings.xlsx, and report.json tests (offline)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

DEMO_DOSSIER = Path(__file__).parent.parent / "fixtures" / "demo_dossier"

PHASE_1_4_ARTIFACTS = (
    "claims.jsonl",
    "citations.jsonl",
    "sources.jsonl",
    "evidence_links.jsonl",
    "findings.jsonl",
    "findings.csv",
    "graph.jsonl",
    "audit_manifest.json",
)

PHASE_5_ARTIFACTS = (
    "report.json",
    "findings.xlsx",
    "report.docx",
)

ALL_EXPECTED_ARTIFACTS = PHASE_1_4_ARTIFACTS + PHASE_5_ARTIFACTS

FORBIDDEN_REPORT_TERMS = (
    "non-compliant",
    "noncompliant",
    "regulatory failure",
    "false claim",
    "nb will reject",
    "notified body will reject",
    "mdr violation",
    "unsupported",
    "must ",
    "shall ",
)


def _assert_no_forbidden(text: str, where: str) -> None:
    lower = text.lower()
    offenders = [term for term in FORBIDDEN_REPORT_TERMS if term in lower]
    assert not offenders, f"Forbidden term(s) {offenders} found in {where}"


@pytest.fixture(scope="module")
def demo_run(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Run verify_dossier once on the demo dossier and cache the output dir."""
    from locuslab.pipeline import verify_dossier

    run_dir = tmp_path_factory.mktemp("phase5_demo_run")
    verify_dossier(DEMO_DOSSIER, run_dir)
    return run_dir


@pytest.fixture(scope="module")
def demo_report_json(demo_run: Path) -> dict[str, object]:
    return json.loads((demo_run / "report.json").read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def demo_audit_manifest(demo_run: Path) -> dict[str, object]:
    return json.loads((demo_run / "audit_manifest.json").read_text(encoding="utf-8"))


class TestArtifactPresence:
    def test_all_eleven_artifacts_written(self, demo_run: Path) -> None:
        for fname in ALL_EXPECTED_ARTIFACTS:
            assert (demo_run / fname).exists(), f"Missing artifact: {fname}"

    def test_report_artifacts_are_non_empty(self, demo_run: Path) -> None:
        for fname in PHASE_5_ARTIFACTS:
            assert (demo_run / fname).stat().st_size > 0, f"Empty: {fname}"


class TestReportJsonShape:
    def test_parses_as_json(self, demo_report_json: dict[str, object]) -> None:
        assert isinstance(demo_report_json, dict)

    def test_required_top_level_keys(self, demo_report_json: dict[str, object]) -> None:
        required = {
            "report_schema_version",
            "run_id",
            "dossier_path",
            "artifact_counts",
            "input_documents",
            "findings_summary",
            "findings_detail",
            "observed_methods",
            "known_limitations",
            "source_artifact_hashes",
        }
        missing = required - demo_report_json.keys()
        assert not missing, f"report.json missing keys: {missing}"

    def test_schema_version_pinned(self, demo_report_json: dict[str, object]) -> None:
        assert demo_report_json["report_schema_version"] == "report.v1"

    def test_run_id_matches_audit_manifest(
        self,
        demo_report_json: dict[str, object],
        demo_audit_manifest: dict[str, object],
    ) -> None:
        assert demo_report_json["run_id"] == demo_audit_manifest["run_id"]

    def test_artifact_counts_findings_seven_on_demo(
        self, demo_report_json: dict[str, object]
    ) -> None:
        counts = demo_report_json["artifact_counts"]
        assert counts["findings"] == 7  # type: ignore[index]

    def test_findings_detail_length_matches_findings_jsonl(
        self, demo_run: Path, demo_report_json: dict[str, object]
    ) -> None:
        findings_jsonl_count = sum(
            1
            for line in (demo_run / "findings.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        )
        assert len(demo_report_json["findings_detail"]) == findings_jsonl_count  # type: ignore[arg-type]

    def test_findings_summary_by_severity_totals_match(
        self, demo_report_json: dict[str, object]
    ) -> None:
        summary = demo_report_json["findings_summary"]
        by_sev = summary["by_severity"]  # type: ignore[index]
        total = sum(by_sev.values())  # type: ignore[union-attr]
        detail_len = len(demo_report_json["findings_detail"])  # type: ignore[arg-type]
        assert total == detail_len

    def test_known_limitations_quote_audit_manifest(
        self,
        demo_report_json: dict[str, object],
        demo_audit_manifest: dict[str, object],
    ) -> None:
        assert demo_report_json["known_limitations"] == demo_audit_manifest["known_limitations"]

    def test_source_artifact_hashes_are_verified_subset_of_manifest(
        self,
        demo_report_json: dict[str, object],
        demo_audit_manifest: dict[str, object],
    ) -> None:
        source_hashes = demo_report_json["source_artifact_hashes"]
        manifest_hashes = demo_audit_manifest["artifact_hashes"]
        assert source_hashes
        assert source_hashes.items() <= manifest_hashes.items()  # type: ignore[union-attr]
        assert not {
            "report.json",
            "findings.xlsx",
            "report.docx",
        } & source_hashes.keys()  # type: ignore[union-attr]


class TestReportJsonForbiddenLanguage:
    """Forbidden-language ban applies to verdict-bearing prose only.

    Engine state names that incidentally contain a banned substring (e.g. the
    parser name `unsupported-file` or the warning code `unsupported_file_type`
    for unreadable extensions in the demo dossier) are legitimate traceability
    and not verdicts. We assert on the fields that actually carry prose:
    known_limitations, finding evidence, finding remediation_hint.
    """

    def test_known_limitations_has_no_forbidden_terms(
        self, demo_report_json: dict[str, object]
    ) -> None:
        for item in demo_report_json["known_limitations"]:  # type: ignore[union-attr]
            _assert_no_forbidden(str(item), "report.json:known_limitations")

    def test_finding_prose_has_no_forbidden_terms(
        self, demo_report_json: dict[str, object]
    ) -> None:
        for f in demo_report_json["findings_detail"]:  # type: ignore[union-attr]
            _assert_no_forbidden(
                str(f["evidence"]),
                f"report.json:findings_detail[{f['eco_id']!r}].evidence",
            )
            _assert_no_forbidden(
                str(f["remediation_hint"]),
                f"report.json:findings_detail[{f['eco_id']!r}].remediation_hint",
            )


class TestReportJsonDeterminism:
    def test_two_runs_produce_byte_equal_report_json(self, tmp_path: Path) -> None:
        from locuslab.pipeline import verify_dossier

        run_a = tmp_path / "run_a"
        run_b = tmp_path / "run_b"
        verify_dossier(DEMO_DOSSIER, run_a)
        verify_dossier(DEMO_DOSSIER, run_b)
        assert (run_a / "report.json").read_bytes() == (run_b / "report.json").read_bytes()


class TestFindingsXlsxShape:
    def test_opens_with_openpyxl(self, demo_run: Path) -> None:
        from openpyxl import load_workbook

        wb = load_workbook(demo_run / "findings.xlsx", read_only=True)
        assert wb is not None

    def test_findings_sheet_exists(self, demo_run: Path) -> None:
        from openpyxl import load_workbook

        wb = load_workbook(demo_run / "findings.xlsx", read_only=True)
        assert "Findings" in wb.sheetnames

    def test_findings_sheet_has_one_row_per_finding(self, demo_run: Path) -> None:
        from openpyxl import load_workbook

        wb = load_workbook(demo_run / "findings.xlsx", read_only=True)
        ws = wb["Findings"]
        rows = list(ws.iter_rows(values_only=True))
        # Row 0 is the header; remaining rows are findings.
        findings_jsonl_count = sum(
            1
            for line in (demo_run / "findings.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        )
        assert len(rows) - 1 == findings_jsonl_count

    def test_findings_sheet_includes_every_eco_id(self, demo_run: Path) -> None:
        from openpyxl import load_workbook

        eco_ids = {
            json.loads(line)["eco_id"]
            for line in (demo_run / "findings.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        }
        wb = load_workbook(demo_run / "findings.xlsx", read_only=True)
        ws = wb["Findings"]
        rows = list(ws.iter_rows(values_only=True))
        # First column is eco_id by spec.
        sheet_eco_ids = {r[0] for r in rows[1:]}
        assert eco_ids.issubset(sheet_eco_ids)

    def test_findings_sheet_includes_adjudication_columns(self, demo_run: Path) -> None:
        from openpyxl import load_workbook

        wb = load_workbook(demo_run / "findings.xlsx", read_only=True)
        ws = wb["Findings"]
        header = next(ws.iter_rows(values_only=True))
        for col in ("adjudication_state", "reviewer", "review_notes", "resolution"):
            assert col in header, f"Findings sheet missing column {col!r}"


class TestReportDocxShape:
    def test_opens_with_python_docx(self, demo_run: Path) -> None:
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        assert doc is not None

    def test_docx_contains_every_eco_id(self, demo_run: Path) -> None:
        import docx

        eco_ids = {
            json.loads(line)["eco_id"]
            for line in (demo_run / "findings.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        }
        doc = docx.Document(str(demo_run / "report.docx"))
        all_text_parts: list[str] = []
        for paragraph in doc.paragraphs:
            all_text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)
        full_text = "\n".join(all_text_parts)
        missing = [eid for eid in eco_ids if eid not in full_text]
        assert not missing, f"ECO IDs missing from report.docx: {missing}"

    def test_docx_contains_run_summary_text(self, demo_run: Path) -> None:
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for needle in ("Run summary", "run_"):
            assert needle in full_text, f"{needle!r} not found in report.docx"

    def test_docx_contains_known_limitations_section(self, demo_run: Path) -> None:
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        headings = [
            p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")
        ]
        assert any("known limitations" in h.lower() for h in headings), (
            f"No 'known limitations' heading found; headings observed: {headings}"
        )

    def test_docx_has_all_expected_h1_headings(self, demo_run: Path) -> None:
        """All 8 spec-mandated Heading 1 sections must appear in report.docx.

        Regression guard against silent section drop or rename. The title
        ('LocusLab Evidence Trace Audit Report') is at Title style, not
        Heading 1, so it is checked separately by the run-summary test.
        """
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        h1_texts = {
            p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 1"
        }
        expected = {
            "Run summary",
            "Input documents",
            "Artifact inventory",
            "Findings summary",
            "Finding detail",
            "Known limitations",
            "Audit and provenance summary",
            "Reviewer next steps",
        }
        missing = expected - h1_texts
        assert not missing, (
            f"Missing Heading 1 sections in report.docx: {missing}. "
            f"Observed Heading 1: {sorted(h1_texts)}"
        )

    def test_docx_finding_detail_table_has_eight_columns(self, demo_run: Path) -> None:
        """Spec §4 item 6 requires affected_object_ids and adjudication_state
        in the finding detail table alongside the other 6 columns."""
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        # The finding detail table is the one whose header row first cell is "eco_id".
        finding_tables = [
            t for t in doc.tables if t.rows and t.rows[0].cells[0].text == "eco_id"
        ]
        assert finding_tables, "No finding detail table found in report.docx"
        # Distinguish from the findings.csv-mirror; pick the table with 8 cols.
        detail_table = next(
            (t for t in finding_tables if len(t.rows[0].cells) == 8), None
        )
        assert detail_table is not None, (
            f"No 8-column finding detail table found; "
            f"observed column counts: {[len(t.rows[0].cells) for t in finding_tables]}"
        )
        header = [c.text for c in detail_table.rows[0].cells]
        for required in ("affected_object_ids", "adjudication_state"):
            assert required in header, (
                f"Finding detail table missing required column {required!r}; "
                f"observed: {header}"
            )


class TestReportDocxForbiddenLanguage:
    """Same scoping rationale as TestReportJsonForbiddenLanguage.

    Check Phase 5-generated PROSE paragraphs (titles, run summary, known
    limitations bullets, reviewer next steps, audit/provenance sentences,
    artifact-inventory body, finding-detail headings) without flagging
    engine-state strings rendered into table cells (parser names, warning
    codes, identifiers).
    """

    def test_docx_paragraphs_have_no_forbidden_terms(self, demo_run: Path) -> None:
        import docx

        doc = docx.Document(str(demo_run / "report.docx"))
        for p in doc.paragraphs:
            _assert_no_forbidden(p.text, f"report.docx paragraph: {p.text[:60]!r}")


class TestPhase4UnchangedByPhase5:
    """Phase 5 must not modify Phase 4 artifacts.

    The byte-equal check on graph.jsonl and audit_manifest.json across two
    runs already exercises Phase 4 determinism; this test re-asserts the
    invariant in the Phase 5 test file so a Phase 5 regression that mutates
    Phase 4 artifacts is caught locally.
    """

    def test_audit_manifest_byte_equal_across_runs(self, tmp_path: Path) -> None:
        from locuslab.pipeline import verify_dossier

        run_a = tmp_path / "run_a"
        run_b = tmp_path / "run_b"
        verify_dossier(DEMO_DOSSIER, run_a)
        verify_dossier(DEMO_DOSSIER, run_b)
        assert (run_a / "audit_manifest.json").read_bytes() == (
            run_b / "audit_manifest.json"
        ).read_bytes()

    def test_graph_jsonl_byte_equal_across_runs(self, tmp_path: Path) -> None:
        from locuslab.pipeline import verify_dossier

        run_a = tmp_path / "run_a"
        run_b = tmp_path / "run_b"
        verify_dossier(DEMO_DOSSIER, run_a)
        verify_dossier(DEMO_DOSSIER, run_b)
        assert (run_a / "graph.jsonl").read_bytes() == (
            run_b / "graph.jsonl"
        ).read_bytes()

    @pytest.mark.parametrize("artifact_name", ["findings.xlsx", "report.docx"])
    def test_ooxml_reports_are_byte_equal_across_runs(
        self, tmp_path: Path, artifact_name: str
    ) -> None:
        from locuslab.pipeline import verify_dossier

        run_a = tmp_path / "run_a"
        run_b = tmp_path / "run_b"
        verify_dossier(DEMO_DOSSIER, run_a)
        verify_dossier(DEMO_DOSSIER, run_b)

        assert (run_a / artifact_name).read_bytes() == (
            run_b / artifact_name
        ).read_bytes()


class TestForbiddenLanguageHelper:
    def test_helper_raises_on_forbidden_term(self) -> None:
        from locuslab.report import assert_no_forbidden_language

        with pytest.raises(ValueError):
            assert_no_forbidden_language("This claim is non-compliant.", "test")

    def test_helper_raises_on_must_with_trailing_space(self) -> None:
        from locuslab.report import assert_no_forbidden_language

        with pytest.raises(ValueError):
            assert_no_forbidden_language("The manufacturer must update this.", "test")

    def test_helper_accepts_clean_text(self) -> None:
        from locuslab.report import assert_no_forbidden_language

        assert_no_forbidden_language(
            "Claim has no resolved local source. Review item requires manual review.",
            "test",
        )

    def test_helper_does_not_match_must_inside_word(self) -> None:
        """`mustang` should NOT match `must ` (trailing-space convention)."""
        from locuslab.report import assert_no_forbidden_language

        assert_no_forbidden_language("Mustang Class IIa device.", "test")
